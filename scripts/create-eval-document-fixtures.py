#!/usr/bin/env python
# Copyright 2026 DataInfra-RedactionEverything Contributors
# SPDX-License-Identifier: Apache-2.0

"""Generate public DOCX/PDF benchmark fixtures.

The fixtures are synthetic and safe to ship or regenerate. They cover
searchable PDF and DOCX text redaction without depending on private files.
"""

from __future__ import annotations

import argparse
import os
import time
from contextlib import contextmanager
from pathlib import Path

import fitz
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT_DIR = ROOT / "fixtures" / "benchmark"


def temp_target(path: Path) -> Path:
    return path.with_name(f".{path.stem}.{os.getpid()}.tmp{path.suffix}")


@contextmanager
def output_dir_lock(output_dir: Path):
    """Serialize fixture regeneration for the same output directory.

    Windows can reject atomic replace when a concurrent quality gate is reading
    or replacing the same DOCX/PDF fixture. A tiny lock file keeps parallel
    local gates deterministic without adding a third-party dependency.
    """
    lock_path = output_dir / ".create-eval-document-fixtures.lock"
    deadline = time.monotonic() + 60.0
    fd: int | None = None
    while fd is None:
        try:
            fd = os.open(lock_path, os.O_CREAT | os.O_EXCL | os.O_RDWR)
            os.write(fd, str(os.getpid()).encode("ascii", errors="ignore"))
        except FileExistsError:
            try:
                age = time.time() - lock_path.stat().st_mtime
                if age > 120:
                    lock_path.unlink(missing_ok=True)
                    continue
            except OSError:
                pass
            if time.monotonic() >= deadline:
                raise TimeoutError(f"timed out waiting for fixture lock: {lock_path}")
            time.sleep(0.1)
    try:
        yield
    finally:
        if fd is not None:
            os.close(fd)
        lock_path.unlink(missing_ok=True)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Create public benchmark DOCX/PDF fixtures.")
    parser.add_argument("output_dir", nargs="?", default=str(DEFAULT_OUTPUT_DIR))
    return parser.parse_args()


def create_docx(path: Path) -> None:
    tmp_path = temp_target(path)
    doc = Document()
    doc.core_properties.title = "DataInfra public redaction fixture"
    doc.add_heading("Synthetic Data Sharing Agreement", level=1)
    doc.add_paragraph("Data owner: Alice Zhang")
    doc.add_paragraph("Mobile: 13451775049")
    doc.add_paragraph("Email: alice.zhang@example.test")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Account"
    table.cell(0, 1).text = "6222020202020202020"
    table.cell(1, 0).text = "Project"
    table.cell(1, 1).text = "OPEN-2026-BENCHMARK"
    doc.sections[0].header.add_paragraph("Header contact: 010-87654321")
    try:
        doc.save(tmp_path)
        tmp_path.replace(path)
    finally:
        tmp_path.unlink(missing_ok=True)


def create_pdf(path: Path) -> None:
    tmp_path = temp_target(path)
    doc = fitz.open()
    page = doc.new_page(width=540, height=320)
    rows = [
        "Synthetic Data Sharing Agreement",
        "Data owner: Alice Zhang",
        "Mobile: 13451775049",
        "Email: alice.zhang@example.test",
        "Account: 6222020202020202020",
        "Project: OPEN-2026-BENCHMARK",
    ]
    y = 62
    for index, row in enumerate(rows):
        page.insert_text(fitz.Point(72, y), row, fontsize=16 if index == 0 else 12)
        y += 34
    try:
        doc.save(tmp_path)
        tmp_path.replace(path)
    finally:
        doc.close()
        tmp_path.unlink(missing_ok=True)


def main() -> int:
    args = parse_args()
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    docx_path = output_dir / "sample-redaction.docx"
    pdf_path = output_dir / "sample-redaction.pdf"
    with output_dir_lock(output_dir):
        create_docx(docx_path)
        create_pdf(pdf_path)
    print(f"document fixtures: {output_dir}")
    print(f"- {docx_path}")
    print(f"- {pdf_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
