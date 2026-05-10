# Copyright 2026 DataInfra-RedactionEverything Contributors
# SPDX-License-Identifier: Apache-2.0

import asyncio
import os
import zipfile

from docx import Document

from app.models.schemas import Entity, ReplacementMode
from app.services.redaction.replacement_strategy import RedactionContext
from app.services.redaction.text_redactor import TextRedactorMixin


class _TextRedactor(TextRedactorMixin):
    pass


def _read_zip_text(path, name: str) -> str:
    with zipfile.ZipFile(path) as zf:
        return zf.read(name).decode("utf-8")


def _all_docx_xml(path) -> str:
    chunks = []
    with zipfile.ZipFile(path) as zf:
        for name in zf.namelist():
            if name.startswith("word/") and name.endswith(".xml"):
                chunks.append(zf.read(name).decode("utf-8"))
    return "\n".join(chunks)


def _write_docx_with_hidden_text_parts(path) -> None:
    doc = Document()
    doc.add_paragraph("Body owner Alice Body")
    doc.sections[0].header.add_paragraph("Header phone 13451775049")
    doc.save(path)

    textbox_xml = """
      <w:p>
        <w:r>
          <w:pict>
            <v:shape id="TextBox1" type="#_x0000_t202" style="width:200pt;height:50pt">
              <v:textbox>
                <w:txbxContent>
                  <w:p>
                    <w:r><w:t>Alice</w:t></w:r>
                    <w:r><w:t> Textbox</w:t></w:r>
                  </w:p>
                </w:txbxContent>
              </v:textbox>
            </v:shape>
          </w:pict>
        </w:r>
      </w:p>
    """
    comments_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:comment w:id="0" w:author="DataInfra" w:date="2026-01-01T00:00:00Z">
    <w:p><w:r><w:t>Alice</w:t></w:r><w:r><w:t> Comment</w:t></w:r></w:p>
  </w:comment>
</w:comments>
"""
    footnotes_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:footnote w:id="1">
    <w:p><w:r><w:t>Alice</w:t></w:r><w:r><w:t> Footnote</w:t></w:r></w:p>
  </w:footnote>
</w:footnotes>
"""

    with zipfile.ZipFile(path) as zf:
        entries = {name: zf.read(name) for name in zf.namelist()}
        document_xml = zf.read("word/document.xml").decode("utf-8")
        content_types = zf.read("[Content_Types].xml").decode("utf-8")
        rels = zf.read("word/_rels/document.xml.rels").decode("utf-8")

    entries["word/document.xml"] = document_xml.replace("</w:body>", f"{textbox_xml}</w:body>").encode("utf-8")
    entries["[Content_Types].xml"] = content_types.replace(
        "</Types>",
        '<Override PartName="/word/comments.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/>'
        '<Override PartName="/word/footnotes.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/>'
        "</Types>",
    ).encode("utf-8")
    entries["word/_rels/document.xml.rels"] = rels.replace(
        "</Relationships>",
        '<Relationship Id="rIdCommentsDataInfra" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" Target="comments.xml"/>'
        '<Relationship Id="rIdFootnotesDataInfra" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" Target="footnotes.xml"/>'
        "</Relationships>",
    ).encode("utf-8")
    entries["word/comments.xml"] = comments_xml.encode("utf-8")
    entries["word/footnotes.xml"] = footnotes_xml.encode("utf-8")

    tmp_path = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for name, data in entries.items():
            zf.writestr(name, data)
    os.replace(tmp_path, path)


def test_docx_redaction_reaches_textboxes_comments_and_footnotes(tmp_path):
    input_path = tmp_path / "source.docx"
    output_path = tmp_path / "redacted.docx"
    _write_docx_with_hidden_text_parts(input_path)

    context = RedactionContext(ReplacementMode.CUSTOM)
    context.set_custom_replacements({
        "Alice Body": "PERSON_BODY",
        "13451775049": "PHONE_001",
        "Alice Textbox": "PERSON_TEXTBOX",
        "Alice Comment": "PERSON_COMMENT",
        "Alice Footnote": "PERSON_FOOTNOTE",
    })
    entities = [
        Entity(id="e1", text="Alice Body", type="PERSON", start=0, end=10),
        Entity(id="e2", text="13451775049", type="PHONE", start=0, end=11),
        Entity(id="e3", text="Alice Textbox", type="PERSON", start=0, end=13),
        Entity(id="e4", text="Alice Comment", type="PERSON", start=0, end=13),
        Entity(id="e5", text="Alice Footnote", type="PERSON", start=0, end=14),
    ]

    redacted_count = asyncio.run(
        _TextRedactor()._redact_docx(str(input_path), str(output_path), entities, context)
    )

    all_xml = _all_docx_xml(output_path)
    assert redacted_count == 5
    assert "Alice Body" not in all_xml
    assert "13451775049" not in all_xml
    assert "Alice Textbox" not in all_xml
    assert "Alice Comment" not in all_xml
    assert "Alice Footnote" not in all_xml
    assert "PERSON_BODY" in all_xml
    assert "PHONE_001" in all_xml
    assert "PERSON_TEXTBOX" in _read_zip_text(output_path, "word/document.xml")
    assert "PERSON_COMMENT" in _read_zip_text(output_path, "word/comments.xml")
    assert "PERSON_FOOTNOTE" in _read_zip_text(output_path, "word/footnotes.xml")
