"""
文本匿名化模块
处理 DOCX、PDF、TXT 文档的文本替换逻辑
"""
import json
import logging
import os
from collections import Counter
from datetime import datetime
from typing import Any

import fitz
from docx import Document
from docx.opc.constants import CONTENT_TYPE as CT
from lxml import etree

from app.core.config import settings
from app.models.schemas import Entity
from app.services.redaction.replacement_strategy import RedactionContext

logger = logging.getLogger(__name__)

WORD_XML_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


class TextRedactorMixin:
    """
    文本匿名化方法集合
    设计为 mixin，由 Redactor 类继承使用
    """

    async def _redact_docx(
        self,
        input_path: str,
        output_path: str,
        entities: list[Entity],
        context: RedactionContext,
    ) -> int:
        """Word 文档匿名化"""
        doc = Document(input_path)
        redacted_count = 0

        # 构建替换映射
        replacements = {}
        for entity in entities:
            if entity.text not in replacements:
                replacements[entity.text] = context.get_replacement(entity)

        trace_enabled = self._is_docx_font_trace_enabled()
        trace_path = self._get_docx_font_trace_path() if trace_enabled else None
        if trace_enabled and trace_path:
            self._init_docx_font_trace(trace_path, input_path, output_path, replacements)

        for para_idx, para in enumerate(self._iter_all_paragraphs(doc)):
            redacted_count += self._replace_in_paragraph(
                para,
                replacements,
                para_idx=para_idx,
                trace_enabled=trace_enabled,
                trace_path=trace_path,
            )

        redacted_count += self._replace_in_docx_xml_parts(doc, replacements)
        doc.save(output_path)
        return redacted_count

    def _replace_in_docx_xml_parts(self, doc: Document, replacements: dict[str, str]) -> int:
        """Replace text in DOCX XML parts not exposed by python-docx objects."""
        if not replacements:
            return 0
        target_content_types = {
            CT.WML_DOCUMENT_MAIN,
            CT.WML_HEADER,
            CT.WML_FOOTER,
            CT.WML_COMMENTS,
            CT.WML_FOOTNOTES,
            CT.WML_ENDNOTES,
        }
        replaced_count = 0
        for part in doc.part.package.parts:
            if getattr(part, "content_type", None) not in target_content_types:
                continue
            root = getattr(part, "element", None)
            if root is None and hasattr(part, "blob"):
                try:
                    root = etree.fromstring(part.blob)
                except (etree.XMLSyntaxError, TypeError, ValueError):
                    root = None
            if root is None:
                continue
            part_replaced_count = 0
            for paragraph in self._docx_xpath(root, ".//w:p"):
                part_replaced_count += self._replace_in_docx_xml_paragraph(paragraph, replacements)
            if part_replaced_count and getattr(part, "element", None) is None:
                part._blob = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            replaced_count += part_replaced_count
        return replaced_count

    def _replace_in_docx_xml_paragraph(self, paragraph, replacements: dict[str, str]) -> int:
        text_nodes = list(self._docx_xpath(paragraph, ".//w:t"))
        if not text_nodes:
            return 0
        full_text = "".join(node.text or "" for node in text_nodes)
        if not full_text:
            return 0

        node_ids: list[int] = []
        for index, node in enumerate(text_nodes):
            node_ids.extend([index] * len(node.text or ""))
        if not node_ids:
            return 0

        matches: list[tuple[int, int, str]] = []
        for old_text, new_text in replacements.items():
            if not old_text:
                continue
            start = 0
            while True:
                pos = full_text.find(old_text, start)
                if pos < 0:
                    break
                matches.append((pos, pos + len(old_text), new_text))
                start = pos + len(old_text)

        if not matches:
            return 0
        matches.sort(key=lambda item: (item[0], -(item[1] - item[0])))

        filtered_matches: list[tuple[int, int, str]] = []
        last_end = -1
        for start, end, replacement in matches:
            if start < last_end:
                continue
            filtered_matches.append((start, end, replacement))
            last_end = end

        replace_map: dict[int, tuple[int, str, int]] = {}
        for start, end, replacement in filtered_matches:
            span_node_ids = node_ids[start:end] if end <= len(node_ids) else node_ids[start:]
            target_node_idx = Counter(span_node_ids).most_common(1)[0][0] if span_node_ids else node_ids[start]
            replace_map[start] = (end, replacement, target_node_idx)

        outputs: list[list[str]] = [[] for _ in text_nodes]
        i = 0
        replaced_count = 0
        while i < len(full_text):
            replacement = replace_map.get(i)
            if replacement:
                end, replacement_text, target_node_idx = replacement
                outputs[target_node_idx].append(replacement_text)
                i = end
                replaced_count += 1
            else:
                node_idx = node_ids[i]
                outputs[node_idx].append(full_text[i])
                i += 1

        for index, node in enumerate(text_nodes):
            node.text = "".join(outputs[index])
        return replaced_count

    @staticmethod
    def _docx_xpath(element, query: str):
        try:
            return element.xpath(query)
        except etree.XPathError:
            return element.xpath(query, namespaces=WORD_XML_NS)

    def _iter_all_paragraphs(self, doc: Document):
        """遍历正文/表格/页眉页脚中的所有段落"""
        for para in doc.paragraphs:
            yield para
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        yield para
        for section in doc.sections:
            for para in section.header.paragraphs:
                yield para
            for para in section.footer.paragraphs:
                yield para
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            yield para
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            yield para

    def _replace_in_paragraph(
        self,
        para,
        replacements: dict[str, str],
        para_idx: int | None = None,
        trace_enabled: bool = False,
        trace_path: str | None = None,
    ) -> int:
        """在段落内进行 run 级替换，尽量保留原始格式"""
        if not replacements:
            return 0
        runs = list(para.runs)
        if not runs:
            return 0

        full_text = "".join(run.text for run in runs)
        if not full_text:
            return 0

        # 记录每个字符所属的 run 索引
        style_ids: list[int] = []
        for idx, run in enumerate(runs):
            style_ids.extend([idx] * len(run.text))

        if not style_ids:
            return 0

        # 找到所有替换
        matches: list[tuple[int, int, str]] = []
        for old_text, new_text in replacements.items():
            if not old_text:
                continue
            start = 0
            while True:
                pos = full_text.find(old_text, start)
                if pos < 0:
                    break
                matches.append((pos, pos + len(old_text), new_text))
                start = pos + len(old_text)

        if not matches:
            return 0

        # 优先长匹配，避免"张三丰"被"张三"提前吞掉
        matches.sort(key=lambda x: (x[0], -(x[1] - x[0])))

        # 过滤重叠匹配
        filtered_matches: list[tuple[int, int, str]] = []
        last_end = -1
        for start, end, replacement in matches:
            if start < last_end:
                continue
            filtered_matches.append((start, end, replacement))
            last_end = end

        if not filtered_matches:
            return 0

        before_snapshot = None
        if trace_enabled and trace_path:
            before_snapshot = self._collect_runs_font_snapshot(runs)

        # 构建替换起点索引：start -> (end, replacement, target_run_idx)
        # target_run_idx 使用区间内主样式 run，避免跨 run 时字体错位
        replace_map: dict[int, tuple[int, str, int]] = {}
        for start, end, replacement in filtered_matches:
            span_style_ids = style_ids[start:end] if end <= len(style_ids) else style_ids[start:]
            if span_style_ids:
                target_run_idx = Counter(span_style_ids).most_common(1)[0][0]
            else:
                target_run_idx = style_ids[start] if start < len(style_ids) else style_ids[-1]
            replace_map[start] = (end, replacement, target_run_idx)

        # 按全局文本顺序重建"各 run 的文本内容"
        run_outputs: list[list[str]] = [[] for _ in runs]
        i = 0
        replaced_count = 0
        while i < len(full_text):
            repl = replace_map.get(i)
            if repl:
                end, replacement, target_run_idx = repl
                run_outputs[target_run_idx].append(replacement)
                i = end
                replaced_count += 1
            else:
                run_idx = style_ids[i]
                run_outputs[run_idx].append(full_text[i])
                i += 1

        # 就地更新 run 文本：不新增/删除 run，最大化保留原始字体与样式继承链
        for idx, run in enumerate(runs):
            new_text = "".join(run_outputs[idx])
            if run.text != new_text:
                run.text = new_text

        if trace_enabled and trace_path:
            after_snapshot = self._collect_runs_font_snapshot(runs)
            self._append_docx_font_trace(
                trace_path,
                {
                    "timestamp": datetime.now().isoformat(),
                    "paragraph_index": para_idx,
                    "paragraph_text_before": full_text,
                    "matches": [
                        {
                            "start": s,
                            "end": e,
                            "original": full_text[s:e],
                            "replacement": rep,
                        }
                        for (s, e, rep) in filtered_matches
                    ],
                    "runs_before": before_snapshot,
                    "runs_after": after_snapshot,
                },
            )

        return replaced_count

    def _is_docx_font_trace_enabled(self) -> bool:
        """是否启用 docx 字体调试导出"""
        raw = os.getenv("DOCX_FONT_TRACE", "0").strip().lower()
        return raw in {"1", "true", "yes", "on"}

    def _get_docx_font_trace_path(self) -> str:
        """获取 docx 字体调试导出文件路径（JSONL）"""
        custom_path = os.getenv("DOCX_FONT_TRACE_PATH", "").strip()
        if custom_path:
            return custom_path
        return os.path.join(settings.DATA_DIR, "docx_font_trace.jsonl")

    def _init_docx_font_trace(
        self,
        trace_path: str,
        input_path: str,
        output_path: str,
        replacements: dict[str, str],
    ) -> None:
        """初始化调试导出文件并写入会话头"""
        try:
            trace_dir = os.path.dirname(trace_path)
            if trace_dir:
                os.makedirs(trace_dir, exist_ok=True)
            session_header = {
                "type": "session",
                "timestamp": datetime.now().isoformat(),
                "input_path": input_path,
                "output_path": output_path,
                "replacement_count": len(replacements),
            }
            with open(trace_path, "a", encoding="utf-8") as f:
                f.write(json.dumps(session_header, ensure_ascii=False) + "\n")
        except (OSError, ValueError, TypeError) as e:
            logger.error("DOCX_TRACE 初始化失败: %s", e)

    def _append_docx_font_trace(self, trace_path: str, record: dict[str, Any]) -> None:
        """追加一条调试记录到 JSONL"""
        try:
            with open(trace_path, "a", encoding="utf-8") as f:
                f.write(json.dumps(record, ensure_ascii=False) + "\n")
        except (OSError, ValueError, TypeError) as e:
            logger.error("DOCX_TRACE 写入失败: %s", e)

    def _collect_runs_font_snapshot(self, runs) -> list[dict[str, Any]]:
        """采集 run 的字体链快照（rPr/rFonts/字号/样式）"""
        from docx.oxml.ns import qn

        result: list[dict[str, Any]] = []
        for idx, run in enumerate(runs):
            r = run._element
            rPr = r.rPr
            rFonts = rPr.find(qn("w:rFonts")) if rPr is not None else None
            sz = rPr.find(qn("w:sz")) if rPr is not None else None
            szCs = rPr.find(qn("w:szCs")) if rPr is not None else None

            result.append(
                {
                    "run_index": idx,
                    "text": run.text,
                    "style_id": getattr(run.style, "style_id", None) if run.style else None,
                    "style_name": getattr(run.style, "name", None) if run.style else None,
                    "font_name_api": run.font.name,
                    "font_size_api_pt": float(run.font.size.pt) if run.font.size else None,
                    "rFonts": {
                        "ascii": rFonts.get(qn("w:ascii")) if rFonts is not None else None,
                        "hAnsi": rFonts.get(qn("w:hAnsi")) if rFonts is not None else None,
                        "eastAsia": rFonts.get(qn("w:eastAsia")) if rFonts is not None else None,
                        "cs": rFonts.get(qn("w:cs")) if rFonts is not None else None,
                        "asciiTheme": rFonts.get(qn("w:asciiTheme")) if rFonts is not None else None,
                        "hAnsiTheme": rFonts.get(qn("w:hAnsiTheme")) if rFonts is not None else None,
                        "eastAsiaTheme": rFonts.get(qn("w:eastAsiaTheme")) if rFonts is not None else None,
                        "csTheme": rFonts.get(qn("w:csTheme")) if rFonts is not None else None,
                        "hint": rFonts.get(qn("w:hint")) if rFonts is not None else None,
                    },
                    "rPr_size": {
                        "w:sz": sz.get(qn("w:val")) if sz is not None else None,
                        "w:szCs": szCs.get(qn("w:val")) if szCs is not None else None,
                    },
                    "rPr_xml": rPr.xml if rPr is not None else None,
                }
            )
        return result

    def _copy_run_format(self, source_run, target_run):
        """复制 run 的格式样式（克隆底层 rPr，避免字体族丢失）"""
        from copy import deepcopy

        source_r = source_run._element
        target_r = target_run._element

        # 保留字符样式引用（有些文档字体通过字符样式继承）
        try:
            if source_run.style is not None:
                target_run.style = source_run.style
        except (AttributeError, ValueError, KeyError):
            pass

        # 移除目标 run 默认生成的 rPr，避免与源格式叠加冲突
        target_rPr = target_r.rPr
        if target_rPr is not None:
            target_r.remove(target_rPr)

        # 直接克隆源 run 的全部字符格式（包含 rFonts/eastAsia/theme/size/color 等）
        source_rPr = source_r.rPr
        if source_rPr is not None:
            target_r.insert(0, deepcopy(source_rPr))

    async def _redact_txt(
        self,
        input_path: str,
        output_path: str,
        entities: list[Entity],
        context: RedactionContext,
    ) -> int:
        """纯文本文件匿名化（.txt, .md, .html, .rtf）— 单次正则替换，O(n) 遍历"""
        import re as _re

        # 读取原文（兼容多种编码）
        content = None
        for enc in ("utf-8", "gbk", "gb2312", "latin-1"):
            try:
                with open(input_path, encoding=enc) as f:
                    content = f.read()
                break
            except (UnicodeDecodeError, ValueError):
                continue
        if content is None:
            with open(input_path, encoding="utf-8", errors="replace") as f:
                content = f.read()

        # 构建替换映射
        replacements: dict[str, str] = {}
        for entity in entities:
            if entity.text and entity.text not in replacements:
                replacements[entity.text] = context.get_replacement(entity)

        if not replacements:
            # 无需替换，直接拷贝
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(content)
            return 0

        # 构建一个联合正则：按长度降序排列，用 | 连接，单次遍历完成所有替换
        # 比逐个 str.replace 更高效（避免多次全文扫描）
        sorted_keys = sorted(replacements.keys(), key=len, reverse=True)
        pattern = _re.compile("|".join(_re.escape(k) for k in sorted_keys))
        redacted_count = 0

        def _replace_match(m: _re.Match) -> str:
            nonlocal redacted_count
            redacted_count += 1
            return replacements[m.group(0)]

        content = pattern.sub(_replace_match, content)

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(content)

        return redacted_count

    async def _redact_pdf_text(
        self,
        input_path: str,
        output_path: str,
        entities: list[Entity],
        context: RedactionContext,
    ) -> int:
        """PDF 文档匿名化（文本型）"""
        doc = fitz.open(input_path)
        redacted_count = 0

        # 构建替换映射
        replacements = {}
        for entity in entities:
            if entity.text not in replacements:
                replacements[entity.text] = context.get_replacement(entity)

        # 对每一页进行处理
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            replacement_inserts: list[tuple[fitz.Rect, str]] = []

            for old_text, new_text in replacements.items():
                # 查找文本位置
                text_instances = page.search_for(old_text)

                for inst in text_instances:
                    # Use a real PDF redaction annotation so original text is
                    # removed from the content stream, not merely covered.
                    rect = fitz.Rect(inst)
                    page.add_redact_annot(rect, fill=(1, 1, 1))
                    replacement_inserts.append((rect, new_text))

                    redacted_count += 1

            if replacement_inserts:
                page.apply_redactions()
                for rect, new_text in replacement_inserts:
                    page.insert_textbox(
                        rect,
                        new_text,
                        fontsize=self._fit_pdf_replacement_font_size(rect, new_text),
                        color=(0, 0, 0),
                        align=fitz.TEXT_ALIGN_LEFT,
                    )

        doc.save(output_path, garbage=4, deflate=True, clean=True)
        doc.close()

        return redacted_count

    @staticmethod
    def _fit_pdf_replacement_font_size(rect: fitz.Rect, text: str) -> float:
        """Choose a conservative font size for inline PDF replacement labels."""
        if not text:
            return 10.0
        base_size = max(6.0, min(10.0, rect.height * 0.78))
        estimated_width = fitz.get_text_length(text, fontsize=base_size)
        if estimated_width <= max(1.0, rect.width):
            return base_size
        return max(5.0, min(base_size, base_size * rect.width / max(1.0, estimated_width)))

    def _extract_docx_text(self, file_path: str) -> str:
        """提取 Word 文档文本（含表格，与 FileParser._parse_docx 结构一致）"""
        doc = Document(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        return "\n".join(paragraphs)

    def _extract_pdf_text(self, file_path: str) -> str:
        """提取 PDF 文档文本"""
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text() + "\n"
        doc.close()
        return text

    def _read_txt(self, file_path: str) -> str:
        """读取纯文本文件（兼容多编码）"""
        for enc in ("utf-8", "gbk", "gb2312", "latin-1"):
            try:
                with open(file_path, encoding=enc) as f:
                    return f.read()
            except (UnicodeDecodeError, ValueError):
                continue
        with open(file_path, encoding="utf-8", errors="replace") as f:
            return f.read()

    def _safe_extract_text(self, file_path: str, ft: str) -> str:
        """安全提取文本，兼容 .doc 和 .docx"""
        if ft == "doc":
            # .doc 文件不能直接用 python-docx 打开
            # 尝试查找转换后的 .docx 文件
            docx_path = file_path.rsplit(".", 1)[0] + ".docx"
            if os.path.exists(docx_path):
                return self._extract_docx_text(docx_path)
            # 尝试从临时目录找
            import glob
            tmp_pattern = os.path.join(os.path.dirname(file_path), "*.docx")
            docx_files = glob.glob(tmp_pattern)
            for f in docx_files:
                if os.path.basename(file_path).rsplit(".", 1)[0] in os.path.basename(f):
                    return self._extract_docx_text(f)
            return "[.doc 文件无法直接提取文本，请查看原始文档]"
        else:
            return self._extract_docx_text(file_path)
