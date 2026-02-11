"""
脱敏执行服务
处理文本和图片的脱敏逻辑
"""
import os
import uuid
import re
import json
from datetime import datetime
from typing import Optional, Any
from collections import Counter
from docx import Document
import fitz

from app.core.config import settings
from app.models.schemas import (
    Entity, 
    BoundingBox, 
    RedactionConfig, 
    ReplacementMode,
    EntityType,
    FileType,
)
from app.services.vision_service import VisionService


class RedactionContext:
    """
    脱敏上下文
    维护实体映射关系，确保同一实体在文档中的一致性
    """
    
    def __init__(self, mode: ReplacementMode):
        self.mode = mode
        self.entity_map: dict[str, str] = {}
        self._coref_map: dict[str, str] = {}
        self.type_counters: dict[str, int] = {}
        self.custom_replacements: dict[str, str] = {}
        self.structured_tag_map: dict[str, str] = {}
    
    def set_custom_replacements(self, replacements: dict[str, str]):
        """设置自定义替换映射"""
        self.custom_replacements = replacements

    def set_structured_mapping(self, mapping: dict[str, list[str]]):
        """设置结构化标签映射（tag -> 原文列表）"""
        for tag, values in mapping.items():
            for value in values:
                if value and value not in self.structured_tag_map:
                    self.structured_tag_map[value] = tag
    
    def get_replacement(self, entity: Entity) -> str:
        """
        获取实体的替换文本
        确保同一实体在整个文档中使用相同的替换
        """
        # 使用 coref_id 作为主键以保持指代一致
        entity_key = entity.coref_id or entity.text
        if entity_key in self._coref_map:
            return self._coref_map[entity_key]
        
        # 根据模式生成替换文本
        if self.mode == ReplacementMode.CUSTOM:
            # 自定义模式：使用预设的替换
            replacement = self.custom_replacements.get(
                entity.text,
                entity.replacement or self._generate_smart_replacement(entity)
            )
        elif self.mode == ReplacementMode.MASK:
            # 掩码模式
            replacement = self._generate_mask_replacement(entity)
        elif self.mode == ReplacementMode.STRUCTURED:
            # 结构化语义标签
            replacement = self._generate_structured_replacement(entity)
        else:
            # 智能模式
            replacement = self._generate_smart_replacement(entity)
        
        self._coref_map[entity_key] = replacement
        if entity.text not in self.entity_map:
            self.entity_map[entity.text] = replacement
        return replacement
    
    def _generate_smart_replacement(self, entity: Entity) -> str:
        """生成智能替换文本"""
        entity_type = entity.type
        type_key = entity_type.value if isinstance(entity_type, EntityType) else str(entity_type)
        
        # 获取计数器
        if type_key not in self.type_counters:
            self.type_counters[type_key] = 0
        self.type_counters[type_key] += 1
        count = self.type_counters[type_key]
        
        # 根据类型生成替换文本
        type_labels = {
            "PERSON": "当事人",
            "ORG": "公司",
            "ID_CARD": "证件号",
            "PHONE": "电话",
            "ADDRESS": "地址",
            "BANK_CARD": "账号",
            "CASE_NUMBER": "案号",
            "DATE": "日期",
            "MONEY": "金额",
            "AMOUNT": "金额",
            "EMAIL": "邮箱",
            "LICENSE_PLATE": "车牌",
            "CONTRACT_NO": "合同编号",
            "CUSTOM": "敏感信息",
        }
        
        label = type_labels.get(type_key, "敏感信息")
        
        # 使用中文数字编号
        chinese_nums = ["零", "一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
        if count <= 10:
            num_str = chinese_nums[count]
        else:
            num_str = str(count)
        
        return f"[{label}{num_str}]"
    
    def _generate_mask_replacement(self, entity: Entity) -> str:
        """生成掩码替换文本"""
        text = entity.text
        length = len(text)
        type_key = entity.type.value if isinstance(entity.type, EntityType) else str(entity.type)
        
        if type_key == "PERSON":
            # 人名：保留姓，其他用 *
            if length >= 2:
                return text[0] + "*" * (length - 1)
            return "*"
        
        elif type_key == "PHONE":
            # 电话：保留前3后4
            if length >= 11:
                return text[:3] + "****" + text[-4:]
            return "*" * length
        
        elif type_key == "ID_CARD":
            # 身份证：保留前6后4
            if length >= 18:
                return text[:6] + "********" + text[-4:]
            return "*" * length
        
        elif type_key == "BANK_CARD":
            # 银行卡：保留后4
            if length >= 16:
                return "*" * (length - 4) + text[-4:]
            return "*" * length
        
        else:
            # 其他：全部用 *
            return "*" * length

    def _generate_structured_replacement(self, entity: Entity) -> str:
        """生成结构化语义标签"""
        type_key = entity.type.value if isinstance(entity.type, EntityType) else str(entity.type)

        if entity.coref_id and entity.coref_id.startswith("<") and entity.coref_id.endswith(">"):
            return entity.coref_id

        if entity.text in self.structured_tag_map:
            return self.structured_tag_map[entity.text]

        template = self._get_tag_template(type_key)
        if template:
            if type_key not in self.type_counters:
                self.type_counters[type_key] = 0
            self.type_counters[type_key] += 1
            index = self.type_counters[type_key]
            return template.replace("{index}", f"{index:03d}")

        structured_map = {
            "PERSON": ("人物", "个人.姓名"),
            "ORG": ("组织", "企业.完整名称"),
            "ADDRESS": ("地点", "办公地址.完整地址"),
            "PHONE": ("电话", "固定电话.号码"),
            "ID_CARD": ("编号", "身份证.号码"),
            "BANK_CARD": ("编号", "银行卡.号码"),
            "CASE_NUMBER": ("编号", "案件编号.号码"),
            "DATE": ("日期/时间", "具体日期.年月日"),
            "MONEY": ("金额", "合同金额.数值"),
            "AMOUNT": ("金额", "合同金额.数值"),
            "EMAIL": ("邮箱", "个人邮箱.地址"),
            "LICENSE_PLATE": ("编号", "车牌.号码"),
            "CONTRACT_NO": ("编号", "合同编号.代码"),
        }

        if type_key not in self.type_counters:
            self.type_counters[type_key] = 0
        self.type_counters[type_key] += 1
        index = self.type_counters[type_key]

        type_name = structured_map.get(type_key)
        if type_name:
            category, path = type_name
            return f"<{category}[{index:03d}].{path}>"

        # 自定义或未知类型兜底
        label = type_key
        return f"<{label}[{index:03d}].完整名称>"

    def _get_tag_template(self, type_key: str) -> Optional[str]:
        try:
            from app.api.entity_types import entity_types_db
            cfg = entity_types_db.get(type_key)
            if cfg and getattr(cfg, "tag_template", None):
                return cfg.tag_template
        except Exception:
            return None
        return None


class Redactor:
    """脱敏执行器"""
    
    def __init__(self):
        self.vision_service = VisionService()
    
    async def redact(
        self,
        file_info: dict,
        entities: list[Entity],
        bounding_boxes: list[BoundingBox],
        config: RedactionConfig,
    ) -> dict:
        """
        执行脱敏操作
        
        Args:
            file_info: 文件信息
            entities: 要脱敏的实体列表
            bounding_boxes: 要脱敏的图片区域列表
            config: 脱敏配置
            
        Returns:
            脱敏结果
        """
        file_type = file_info["file_type"]
        file_path = file_info["file_path"]
        
        # 创建脱敏上下文
        context = RedactionContext(config.replacement_mode)
        context.set_custom_replacements(config.custom_replacements)
        
        # 生成输出文件路径
        output_file_id = str(uuid.uuid4())
        original_ext = os.path.splitext(file_path)[1]
        output_ext = original_ext
        if file_type == FileType.DOC:
            output_ext = ".docx"
        output_path = os.path.join(settings.OUTPUT_DIR, f"{output_file_id}{output_ext}")
        
        # 只处理选中的实体
        selected_entities = [e for e in entities if e.selected]
        selected_boxes = [b for b in bounding_boxes if b.selected]

        # 结构化模式：优先使用 HaS Hide 映射提升一致性
        if config.replacement_mode == ReplacementMode.STRUCTURED and file_info.get("content"):
            try:
                from app.services.has_service import has_service
                from app.api.entity_types import entity_types_db

                if has_service.is_available():
                    type_ids = {e.type for e in selected_entities}
                    entity_types = [entity_types_db[tid] for tid in type_ids if tid in entity_types_db]
                    if entity_types:
                        masked_text, mapping = await has_service.hide_text(
                            file_info["content"],
                            entity_types,
                        )
                        context.set_structured_mapping(mapping)
                        # 用于前端对比展示
                        file_info["redacted_text"] = masked_text
            except Exception as e:
                print(f"结构化脱敏映射构建失败: {e}")
        
        redacted_count = 0
        
        if file_type == FileType.DOC:
            # 先将 .doc 转换为 .docx 再处理
            converted_path = await self._convert_doc_to_docx(file_path)
            if not converted_path or not os.path.exists(converted_path):
                raise ValueError("DOC 转换失败，无法脱敏")
            redacted_count = await self._redact_docx(
                converted_path, output_path, selected_entities, context
            )
            # 清理转换后的临时文件
            if converted_path != file_path:
                try:
                    os.remove(converted_path)
                except:
                    pass
        elif file_type == FileType.DOCX:
            # Word 文档脱敏
            redacted_count = await self._redact_docx(
                file_path, output_path, selected_entities, context
            )
        elif file_type == FileType.PDF:
            # PDF 文档脱敏（文本型）
            redacted_count = await self._redact_pdf_text(
                file_path, output_path, selected_entities, context
            )
        elif file_type in [FileType.PDF_SCANNED, FileType.IMAGE]:
            # 图片/扫描件脱敏
            await self.vision_service.apply_redaction(
                file_path, file_type, selected_boxes, output_path
            )
            redacted_count = len(selected_boxes)
        
        return {
            "output_file_id": output_file_id,
            "output_path": output_path,
            "redacted_count": redacted_count,
            "entity_map": context.entity_map,
        }

    async def _convert_doc_to_docx(self, file_path: str) -> Optional[str]:
        """将 .doc 转换为 .docx（复用 FileParser 逻辑）"""
        try:
            from app.services.file_parser import FileParser
            parser = FileParser()
            return await parser._convert_doc_to_docx(file_path)
        except Exception as e:
            print(f"DOC 转换失败: {e}")
            return None
    
    async def _redact_docx(
        self,
        input_path: str,
        output_path: str,
        entities: list[Entity],
        context: RedactionContext,
    ) -> int:
        """Word 文档脱敏"""
        doc = Document(input_path)
        redacted_count = 0
        
        # 构建替换映射
        replacements = {}
        for entity in entities:
            if entity.text not in replacements:
                replacements[entity.text] = context.get_replacement(entity)

        trace_enabled = self._is_docx_font_trace_enabled()
        trace_path = self._get_docx_font_trace_path() if trace_enabled else None
        if trace_enabled and trace_path:
            self._init_docx_font_trace(trace_path, input_path, output_path, replacements)

        for para_idx, para in enumerate(self._iter_all_paragraphs(doc)):
            redacted_count += self._replace_in_paragraph(
                para,
                replacements,
                para_idx=para_idx,
                trace_enabled=trace_enabled,
                trace_path=trace_path,
            )
        
        doc.save(output_path)
        return redacted_count

    def _iter_all_paragraphs(self, doc: Document):
        """遍历正文/表格/页眉页脚中的所有段落"""
        for para in doc.paragraphs:
            yield para
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        yield para
        for section in doc.sections:
            for para in section.header.paragraphs:
                yield para
            for para in section.footer.paragraphs:
                yield para
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            yield para
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            yield para

    def _replace_in_paragraph(
        self,
        para,
        replacements: dict[str, str],
        para_idx: Optional[int] = None,
        trace_enabled: bool = False,
        trace_path: Optional[str] = None,
    ) -> int:
        """在段落内进行 run 级替换，尽量保留原始格式"""
        if not replacements:
            return 0
        runs = list(para.runs)
        if not runs:
            return 0

        full_text = "".join(run.text for run in runs)
        if not full_text:
            return 0

        # 记录每个字符所属的 run 索引
        style_ids: list[int] = []
        for idx, run in enumerate(runs):
            style_ids.extend([idx] * len(run.text))

        if not style_ids:
            return 0

        # 找到所有替换
        matches: list[tuple[int, int, str]] = []
        for old_text, new_text in replacements.items():
            if not old_text:
                continue
            start = 0
            while True:
                pos = full_text.find(old_text, start)
                if pos < 0:
                    break
                matches.append((pos, pos + len(old_text), new_text))
                start = pos + len(old_text)

        if not matches:
            return 0

        # 优先长匹配，避免“张三丰”被“张三”提前吞掉
        matches.sort(key=lambda x: (x[0], -(x[1] - x[0])))

        # 过滤重叠匹配
        filtered_matches: list[tuple[int, int, str]] = []
        last_end = -1
        for start, end, replacement in matches:
            if start < last_end:
                continue
            filtered_matches.append((start, end, replacement))
            last_end = end

        if not filtered_matches:
            return 0

        before_snapshot = None
        if trace_enabled and trace_path:
            before_snapshot = self._collect_runs_font_snapshot(runs)

        # 构建替换起点索引：start -> (end, replacement, target_run_idx)
        # target_run_idx 使用区间内主样式 run，避免跨 run 时字体错位
        replace_map: dict[int, tuple[int, str, int]] = {}
        for start, end, replacement in filtered_matches:
            span_style_ids = style_ids[start:end] if end <= len(style_ids) else style_ids[start:]
            if span_style_ids:
                target_run_idx = Counter(span_style_ids).most_common(1)[0][0]
            else:
                target_run_idx = style_ids[start] if start < len(style_ids) else style_ids[-1]
            replace_map[start] = (end, replacement, target_run_idx)

        # 按全局文本顺序重建“各 run 的文本内容”
        run_outputs: list[list[str]] = [[] for _ in runs]
        i = 0
        replaced_count = 0
        while i < len(full_text):
            repl = replace_map.get(i)
            if repl:
                end, replacement, target_run_idx = repl
                run_outputs[target_run_idx].append(replacement)
                i = end
                replaced_count += 1
            else:
                run_idx = style_ids[i]
                run_outputs[run_idx].append(full_text[i])
                i += 1

        # 就地更新 run 文本：不新增/删除 run，最大化保留原始字体与样式继承链
        for idx, run in enumerate(runs):
            new_text = "".join(run_outputs[idx])
            if run.text != new_text:
                run.text = new_text

        if trace_enabled and trace_path:
            after_snapshot = self._collect_runs_font_snapshot(runs)
            self._append_docx_font_trace(
                trace_path,
                {
                    "timestamp": datetime.now().isoformat(),
                    "paragraph_index": para_idx,
                    "paragraph_text_before": full_text,
                    "matches": [
                        {
                            "start": s,
                            "end": e,
                            "original": full_text[s:e],
                            "replacement": rep,
                        }
                        for (s, e, rep) in filtered_matches
                    ],
                    "runs_before": before_snapshot,
                    "runs_after": after_snapshot,
                },
            )

        return replaced_count

    def _is_docx_font_trace_enabled(self) -> bool:
        """是否启用 docx 字体调试导出"""
        raw = os.getenv("DOCX_FONT_TRACE", "0").strip().lower()
        return raw in {"1", "true", "yes", "on"}

    def _get_docx_font_trace_path(self) -> str:
        """获取 docx 字体调试导出文件路径（JSONL）"""
        custom_path = os.getenv("DOCX_FONT_TRACE_PATH", "").strip()
        if custom_path:
            return custom_path
        return os.path.join(settings.DATA_DIR, "docx_font_trace.jsonl")

    def _init_docx_font_trace(
        self,
        trace_path: str,
        input_path: str,
        output_path: str,
        replacements: dict[str, str],
    ) -> None:
        """初始化调试导出文件并写入会话头"""
        try:
            trace_dir = os.path.dirname(trace_path)
            if trace_dir:
                os.makedirs(trace_dir, exist_ok=True)
            session_header = {
                "type": "session",
                "timestamp": datetime.now().isoformat(),
                "input_path": input_path,
                "output_path": output_path,
                "replacement_count": len(replacements),
            }
            with open(trace_path, "a", encoding="utf-8") as f:
                f.write(json.dumps(session_header, ensure_ascii=False) + "\n")
        except Exception as e:
            print(f"[DOCX_TRACE] 初始化失败: {e}")

    def _append_docx_font_trace(self, trace_path: str, record: dict[str, Any]) -> None:
        """追加一条调试记录到 JSONL"""
        try:
            with open(trace_path, "a", encoding="utf-8") as f:
                f.write(json.dumps(record, ensure_ascii=False) + "\n")
        except Exception as e:
            print(f"[DOCX_TRACE] 写入失败: {e}")

    def _collect_runs_font_snapshot(self, runs) -> list[dict[str, Any]]:
        """采集 run 的字体链快照（rPr/rFonts/字号/样式）"""
        from docx.oxml.ns import qn

        result: list[dict[str, Any]] = []
        for idx, run in enumerate(runs):
            r = run._element
            rPr = r.rPr
            rFonts = rPr.find(qn("w:rFonts")) if rPr is not None else None
            sz = rPr.find(qn("w:sz")) if rPr is not None else None
            szCs = rPr.find(qn("w:szCs")) if rPr is not None else None

            result.append(
                {
                    "run_index": idx,
                    "text": run.text,
                    "style_id": getattr(run.style, "style_id", None) if run.style else None,
                    "style_name": getattr(run.style, "name", None) if run.style else None,
                    "font_name_api": run.font.name,
                    "font_size_api_pt": float(run.font.size.pt) if run.font.size else None,
                    "rFonts": {
                        "ascii": rFonts.get(qn("w:ascii")) if rFonts is not None else None,
                        "hAnsi": rFonts.get(qn("w:hAnsi")) if rFonts is not None else None,
                        "eastAsia": rFonts.get(qn("w:eastAsia")) if rFonts is not None else None,
                        "cs": rFonts.get(qn("w:cs")) if rFonts is not None else None,
                        "asciiTheme": rFonts.get(qn("w:asciiTheme")) if rFonts is not None else None,
                        "hAnsiTheme": rFonts.get(qn("w:hAnsiTheme")) if rFonts is not None else None,
                        "eastAsiaTheme": rFonts.get(qn("w:eastAsiaTheme")) if rFonts is not None else None,
                        "csTheme": rFonts.get(qn("w:csTheme")) if rFonts is not None else None,
                        "hint": rFonts.get(qn("w:hint")) if rFonts is not None else None,
                    },
                    "rPr_size": {
                        "w:sz": sz.get(qn("w:val")) if sz is not None else None,
                        "w:szCs": szCs.get(qn("w:val")) if szCs is not None else None,
                    },
                    "rPr_xml": rPr.xml if rPr is not None else None,
                }
            )
        return result

    def _copy_run_format(self, source_run, target_run):
        """复制 run 的格式样式（克隆底层 rPr，避免字体族丢失）"""
        from copy import deepcopy

        source_r = source_run._element
        target_r = target_run._element

        # 保留字符样式引用（有些文档字体通过字符样式继承）
        try:
            if source_run.style is not None:
                target_run.style = source_run.style
        except Exception:
            pass

        # 移除目标 run 默认生成的 rPr，避免与源格式叠加冲突
        target_rPr = target_r.rPr
        if target_rPr is not None:
            target_r.remove(target_rPr)

        # 直接克隆源 run 的全部字符格式（包含 rFonts/eastAsia/theme/size/color 等）
        source_rPr = source_r.rPr
        if source_rPr is not None:
            target_r.insert(0, deepcopy(source_rPr))
    
    async def _redact_pdf_text(
        self,
        input_path: str,
        output_path: str,
        entities: list[Entity],
        context: RedactionContext,
    ) -> int:
        """PDF 文档脱敏（文本型）"""
        doc = fitz.open(input_path)
        redacted_count = 0
        
        # 构建替换映射
        replacements = {}
        for entity in entities:
            if entity.text not in replacements:
                replacements[entity.text] = context.get_replacement(entity)
        
        # 对每一页进行处理
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            
            for old_text, new_text in replacements.items():
                # 查找文本位置
                text_instances = page.search_for(old_text)
                
                for inst in text_instances:
                    # 添加遮罩（白色背景 + 新文本）
                    # 首先用白色矩形覆盖原文本
                    shape = page.new_shape()
                    shape.draw_rect(inst)
                    shape.finish(color=(1, 1, 1), fill=(1, 1, 1))
                    shape.commit()
                    
                    # 然后插入新文本
                    # 计算文本位置（矩形左上角）
                    text_point = fitz.Point(inst.x0, inst.y1 - 2)
                    page.insert_text(
                        text_point,
                        new_text,
                        fontsize=10,
                        color=(0, 0, 0),
                    )
                    
                    redacted_count += 1
        
        doc.save(output_path)
        doc.close()
        
        return redacted_count
    
    async def get_comparison(self, file_info: dict) -> dict:
        """
        获取脱敏前后对比数据
        """
        file_type = file_info["file_type"]
        original_path = file_info["file_path"]
        redacted_path = file_info.get("output_path")
        redacted_text = file_info.get("redacted_text")
        
        if not redacted_path or not os.path.exists(redacted_path):
            raise ValueError("脱敏文件不存在")
        
        # 统一转为字符串比较（兼容枚举和字符串）
        ft = str(file_type.value) if hasattr(file_type, 'value') else str(file_type)
        is_docx = ft in ("docx", "doc")
        is_pdf = ft in ("pdf", "pdf_scanned")
        
        print(f"[Compare] file_type={file_type}, ft={ft}, is_docx={is_docx}, is_pdf={is_pdf}, redacted_path={redacted_path}")
        
        original_content = ""
        redacted_content = ""
        
        # .doc 文件脱敏后输出为 .docx，原始内容从解析缓存读取
        original_content_cached = file_info.get("content", "")
        
        if redacted_text:
            # 使用结构化脱敏文本（更符合展示）
            redacted_content = redacted_text
            if is_docx:
                original_content = original_content_cached or self._safe_extract_text(original_path, ft)
            elif is_pdf:
                original_content = original_content_cached or self._extract_pdf_text(original_path)
            else:
                original_content = "[图片文件，请查看预览]"
        elif is_docx:
            # Word 文档（.doc 和 .docx）
            original_content = original_content_cached or self._safe_extract_text(original_path, ft)
            # 脱敏后的文件一定是 .docx 格式
            redacted_content = self._extract_docx_text(redacted_path)
        elif is_pdf:
            # PDF 文档
            original_content = original_content_cached or self._extract_pdf_text(original_path)
            redacted_content = self._extract_pdf_text(redacted_path)
        else:
            # 图片类：返回提示信息
            original_content = "[图片文件，请查看预览]"
            redacted_content = "[已脱敏图片，请查看预览]"
        
        # 计算变更
        changes = self._compute_changes(
            original_content,
            redacted_content,
            file_info.get("entity_map", {}),
        )
        
        return {
            "original": original_content,
            "redacted": redacted_content,
            "changes": changes,
        }
    
    def _safe_extract_text(self, file_path: str, ft: str) -> str:
        """安全提取文本，兼容 .doc 和 .docx"""
        if ft == "doc":
            # .doc 文件不能直接用 python-docx 打开
            # 尝试查找转换后的 .docx 文件
            docx_path = file_path.rsplit(".", 1)[0] + ".docx"
            if os.path.exists(docx_path):
                return self._extract_docx_text(docx_path)
            # 尝试从临时目录找
            import glob
            tmp_pattern = os.path.join(os.path.dirname(file_path), "*.docx")
            docx_files = glob.glob(tmp_pattern)
            for f in docx_files:
                if os.path.basename(file_path).rsplit(".", 1)[0] in os.path.basename(f):
                    return self._extract_docx_text(f)
            return "[.doc 文件无法直接提取文本，请查看原始文档]"
        else:
            return self._extract_docx_text(file_path)
    
    def _extract_docx_text(self, file_path: str) -> str:
        """提取 Word 文档文本"""
        doc = Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        return "\n".join(paragraphs)
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """提取 PDF 文档文本"""
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text() + "\n"
        doc.close()
        return text
    
    def _compute_changes(
        self,
        original: str,
        redacted: str,
        entity_map: dict[str, str],
    ) -> list[dict]:
        """计算变更列表"""
        changes = []
        
        for original_text, replacement in entity_map.items():
            # 计算出现次数
            count = original.count(original_text)
            if count > 0:
                changes.append({
                    "original": original_text,
                    "replacement": replacement,
                    "count": count,
                })
        
        return changes
