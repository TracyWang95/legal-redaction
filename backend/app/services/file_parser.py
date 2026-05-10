"""
文件解析服务
支持 Word/PDF/图片 文件的内容提取
"""
import logging
import os
import sys
from collections import OrderedDict
from threading import Lock

logger = logging.getLogger(__name__)

import fitz  # PyMuPDF
from docx import Document
from PIL import Image

from app.core.config import settings
from app.core.file_validation import safe_path_in_dir
from app.models.schemas import FileType, ParseResult


def _validate_path(file_path: str) -> None:
    """Raise ``ValueError`` if *file_path* is outside the configured upload/output directories."""
    resolved = os.path.realpath(file_path)
    if safe_path_in_dir(resolved, settings.UPLOAD_DIR):
        return
    if safe_path_in_dir(resolved, settings.OUTPUT_DIR):
        return
    raise ValueError(f"路径不在允许的目录中: {file_path}")


class FileParser:
    """文件解析器"""

    # 判断 PDF 是否为扫描件的文本密度阈值
    TEXT_DENSITY_THRESHOLD = 100  # 每页至少 100 个字符才认为是文本 PDF

    _pdf_page_image_cache: OrderedDict[tuple[str, int, int, int, int], bytes] = OrderedDict()
    _pdf_page_image_cache_lock = Lock()
    _pdf_page_text_blocks_cache: OrderedDict[
        tuple[str, int, int, int, int],
        tuple[list["OCRTextBlock"], int, int],
    ] = OrderedDict()
    _pdf_page_text_blocks_cache_lock = Lock()

    def __init__(self) -> None:
        self.last_pdf_page_image_cache_hit: bool | None = None
        self.last_pdf_page_text_blocks_cache_hit: bool | None = None

    @staticmethod
    def _pdf_page_cache_key(file_path: str, page: int, dpi: int) -> tuple[str, int, int, int, int]:
        resolved = os.path.realpath(file_path)
        try:
            stat = os.stat(resolved)
        except OSError:
            stat = os.stat(file_path)
        return (resolved, int(stat.st_mtime_ns), int(stat.st_size), int(page), int(dpi))

    async def parse(self, file_path: str, file_type: FileType) -> ParseResult:
        """
        解析文件内容

        Args:
            file_path: 文件路径
            file_type: 文件类型

        Returns:
            解析结果

        Raises:
            ValueError: 路径不在允许的目录中
        """
        _validate_path(file_path)
        if file_type == FileType.DOC:
            return await self._parse_doc(file_path)
        elif file_type == FileType.DOCX:
            return await self._parse_docx(file_path)
        elif file_type == FileType.TXT:
            return await self._parse_txt(file_path)
        elif file_type == FileType.PDF:
            return await self._parse_pdf(file_path)
        elif file_type == FileType.IMAGE:
            return await self._parse_image(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {file_type}")

    async def _parse_doc(self, file_path: str) -> ParseResult:
        """
        解析旧版 Word 文档 (.doc)
        """
        # 先尝试转换为 docx
        docx_path = await self._convert_doc_to_docx(file_path)

        if docx_path and os.path.exists(docx_path):
            try:
                result = await self._parse_docx(docx_path)
                result.file_type = FileType.DOC
                return result
            except Exception as e:
                logger.error("解析转换后的 docx 失败: %s", e)
            finally:
                # 清理临时文件
                if docx_path != file_path:
                    try:
                        os.remove(docx_path)
                    except OSError:
                        pass

        # 转换失败，返回提示
        return ParseResult(
            file_id="",
            file_type=FileType.DOC,
            content="[无法解析 .doc 文件，请将文件另存为 .docx 格式后重试]",
            page_count=1,
            pages=[],
            is_scanned=False,
        )

    async def _convert_doc_to_docx(self, doc_path: str) -> str | None:
        """
        将 .doc 文件转换为 .docx
        """
        try:
            _validate_path(doc_path)
        except ValueError:
            logger.warning("Path traversal blocked in _convert_doc_to_docx: %s", doc_path)
            return None
        abs_path = os.path.abspath(doc_path)

        # Windows: 尝试使用 pywin32 调用 Word
        if sys.platform == 'win32':
            result = await self._convert_with_word_com(abs_path)
            if result:
                return result

        # 尝试使用 LibreOffice
        result = await self._convert_with_libreoffice(abs_path)
        if result:
            return result

        return None

    async def _convert_with_word_com(self, doc_path: str) -> str | None:
        """使用 Word COM 接口转换"""
        try:
            _validate_path(doc_path)
        except ValueError:
            logger.warning("Path traversal blocked in _convert_with_word_com: %s", doc_path)
            return None
        try:
            import pythoncom
            import win32com.client

            # 在新线程中需要初始化 COM
            pythoncom.CoInitialize()

            word = None
            doc = None

            try:
                # 创建 Word 应用
                word = win32com.client.DispatchEx("Word.Application")
                word.Visible = False
                word.DisplayAlerts = 0  # wdAlertsNone

                # 打开文档
                doc = word.Documents.Open(
                    doc_path,
                    ReadOnly=True,
                    AddToRecentFiles=False,
                )

                # 生成输出路径
                output_path = doc_path.rsplit('.', 1)[0] + '_converted.docx'

                # 保存为 docx (FileFormat=16 是 docx 格式)
                doc.SaveAs2(output_path, FileFormat=16)

                logger.info("Word COM 转换成功: %s", output_path)
                return output_path

            except Exception as e:
                logger.error("Word COM 转换错误: %s", e)
                return None
            finally:
                # 关闭文档和 Word
                if doc:
                    try:
                        doc.Close(SaveChanges=False)
                    except (OSError, AttributeError, Exception):
                        logger.debug("Failed to close Word document handle")
                if word:
                    try:
                        word.Quit()
                    except (OSError, AttributeError, Exception):
                        logger.debug("Failed to quit Word COM instance")
                pythoncom.CoUninitialize()

        except ImportError:
            logger.warning("pywin32 未安装，无法使用 Word COM")
            return None
        except Exception as e:
            logger.error("Word COM 初始化失败: %s", e)
            return None

    async def _convert_with_libreoffice(self, doc_path: str) -> str | None:
        """使用 LibreOffice 转换"""
        try:
            _validate_path(doc_path)
        except ValueError:
            logger.warning("Path traversal blocked in _convert_with_libreoffice: %s", doc_path)
            return None
        import subprocess

        # 查找 LibreOffice
        soffice_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            "/usr/bin/soffice",
            "/usr/bin/libreoffice",
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        ]

        soffice = None
        for path in soffice_paths:
            if os.path.exists(path):
                soffice = path
                break

        if not soffice:
            logger.warning("未找到 LibreOffice")
            return None

        try:
            output_dir = os.path.dirname(doc_path)

            result = subprocess.run(
                [soffice, "--headless", "--convert-to", "docx", "--outdir", output_dir, doc_path],
                capture_output=True,
                text=True,
                timeout=60,
            )

            if result.returncode == 0:
                output_path = os.path.splitext(doc_path)[0] + '.docx'
                if os.path.exists(output_path):
                    logger.info("LibreOffice 转换成功: %s", output_path)
                    return output_path

            logger.error("LibreOffice 转换失败: %s", result.stderr)
            return None

        except Exception as e:
            logger.error("LibreOffice 转换错误: %s", e)
            return None

    async def _parse_txt(self, file_path: str) -> ParseResult:
        """解析纯文本文件 (.txt, .md, .html, .htm, .rtf)"""
        _validate_path(file_path)
        ext = os.path.splitext(file_path)[1].lower()
        try:
            # 尝试多种编码
            content = None
            for enc in ("utf-8", "gbk", "gb2312", "latin-1"):
                try:
                    with open(file_path, encoding=enc) as f:
                        content = f.read()
                    break
                except (UnicodeDecodeError, ValueError):
                    continue
            if content is None:
                with open(file_path, encoding="utf-8", errors="replace") as f:
                    content = f.read()

            # HTML: 用简单正则去标签提取文本
            if ext in (".html", ".htm"):
                import re
                content = re.sub(r"<style[^>]*>.*?</style>", "", content, flags=re.DOTALL | re.IGNORECASE)
                content = re.sub(r"<script[^>]*>.*?</script>", "", content, flags=re.DOTALL | re.IGNORECASE)
                content = re.sub(r"<[^>]+>", " ", content)
                content = re.sub(r"\s+", " ", content).strip()

            # RTF: 去除 RTF 控制码提取文本
            if ext == ".rtf":
                import re
                content = re.sub(r"\\[a-z]+\d*\s?", "", content)
                content = re.sub(r"[{}]", "", content)
                content = re.sub(r"\s+", " ", content).strip()

            pages = [content] if content else []
            return ParseResult(
                file_id="",
                file_type=FileType.TXT,
                content=content or "",
                page_count=1,
                pages=pages,
                is_scanned=False,
            )
        except Exception as e:
            logger.error("解析文本文件失败: %s", e)
            return ParseResult(
                file_id="",
                file_type=FileType.TXT,
                content=f"[无法解析文件: {os.path.basename(file_path)}]",
                page_count=1,
                pages=[],
                is_scanned=False,
            )

    async def _parse_docx(self, file_path: str) -> ParseResult:
        """解析 Word 文档 (.docx)"""
        _validate_path(file_path)
        doc = Document(file_path)

        paragraphs = []

        # 提取段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # 提取表格
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    paragraphs.append(" | ".join(row_text))

        content = "\n".join(paragraphs)

        return ParseResult(
            file_id="",
            file_type=FileType.DOCX,
            content=content,
            page_count=1,
            pages=[content],
            is_scanned=False,
        )

    async def _parse_pdf(self, file_path: str) -> ParseResult:
        """解析 PDF 文档"""
        _validate_path(file_path)
        doc = fitz.open(file_path)

        pages = []
        total_chars = 0

        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text()
            pages.append(text)
            total_chars += len(text.strip())

        doc.close()

        # 判断是否为扫描件
        avg_chars_per_page = total_chars / len(pages) if pages else 0
        is_scanned = avg_chars_per_page < self.TEXT_DENSITY_THRESHOLD

        content = "\n\n".join(pages)

        return ParseResult(
            file_id="",
            file_type=FileType.PDF_SCANNED if is_scanned else FileType.PDF,
            content=content if not is_scanned else "",
            page_count=len(pages),
            pages=pages if not is_scanned else [],
            is_scanned=is_scanned,
        )

    async def _parse_image(self, file_path: str) -> ParseResult:
        """解析图片文件"""
        return ParseResult(
            file_id="",
            file_type=FileType.IMAGE,
            content="",
            page_count=1,
            pages=[],
            is_scanned=True,
        )

    async def pdf_to_images(self, file_path: str, dpi: int = 150) -> list[bytes]:
        """将 PDF 转换为图片列表"""
        _validate_path(file_path)
        doc = fitz.open(file_path)
        images = []

        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            zoom = dpi / 72
            matrix = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=matrix)
            img_data = pix.tobytes("png")
            images.append(img_data)

        doc.close()
        return images

    async def get_pdf_page_image(self, file_path: str, page: int, dpi: int = 150) -> bytes:
        """获取 PDF 指定页的图片"""
        _validate_path(file_path)
        self.last_pdf_page_image_cache_hit = False
        cache_key = self._pdf_page_cache_key(file_path, page, dpi)
        resolved = cache_key[0]
        cache_limit = int(settings.PDF_PAGE_IMAGE_CACHE_PAGES)
        if cache_limit > 0:
            with self._pdf_page_image_cache_lock:
                cached = self._pdf_page_image_cache.get(cache_key)
                if cached is not None:
                    self._pdf_page_image_cache.move_to_end(cache_key)
                    logger.debug(
                        "PDF page image cache hit: %s page=%d dpi=%d",
                        os.path.basename(file_path),
                        page,
                        dpi,
                    )
                    self.last_pdf_page_image_cache_hit = True
                    return cached

        doc = fitz.open(resolved)

        if page < 1 or page > len(doc):
            doc.close()
            raise ValueError(f"页码超出范围: {page}")

        pdf_page = doc.load_page(page - 1)
        zoom = dpi / 72
        matrix = fitz.Matrix(zoom, zoom)
        pix = pdf_page.get_pixmap(matrix=matrix)
        img_data = pix.tobytes("png")

        doc.close()
        if cache_limit > 0:
            with self._pdf_page_image_cache_lock:
                self._pdf_page_image_cache[cache_key] = img_data
                self._pdf_page_image_cache.move_to_end(cache_key)
                while len(self._pdf_page_image_cache) > cache_limit:
                    self._pdf_page_image_cache.popitem(last=False)
        return img_data

    async def get_pdf_page_text_blocks(self, file_path: str, page: int, dpi: int = 150):
        """Return native PDF text-layer blocks mapped into rendered-page pixels."""
        _validate_path(file_path)
        from app.services.hybrid_vision_service import OCRTextBlock

        self.last_pdf_page_text_blocks_cache_hit = False
        cache_key = self._pdf_page_cache_key(file_path, page, dpi)
        resolved = cache_key[0]
        cache_limit = int(settings.PDF_PAGE_TEXT_BLOCK_CACHE_PAGES)

        def clone_blocks(blocks: list[OCRTextBlock]) -> list[OCRTextBlock]:
            return [
                OCRTextBlock(
                    text=block.text,
                    polygon=[[float(point[0]), float(point[1])] for point in block.polygon],
                    confidence=float(block.confidence),
                )
                for block in blocks
            ]

        if cache_limit > 0:
            with self._pdf_page_text_blocks_cache_lock:
                cached = self._pdf_page_text_blocks_cache.get(cache_key)
                if cached is not None:
                    self._pdf_page_text_blocks_cache.move_to_end(cache_key)
                    blocks, page_width, page_height = cached
                    self.last_pdf_page_text_blocks_cache_hit = True
                    logger.debug(
                        "PDF text-layer cache hit: %s page=%d dpi=%d",
                        os.path.basename(file_path),
                        page,
                        dpi,
                    )
                    return clone_blocks(blocks), page_width, page_height

        doc = fitz.open(resolved)
        try:
            if page < 1 or page > len(doc):
                raise ValueError(f"椤电爜瓒呭嚭鑼冨洿: {page}")

            pdf_page = doc.load_page(page - 1)
            zoom = dpi / 72
            page_width = max(1, int(round(pdf_page.rect.width * zoom)))
            page_height = max(1, int(round(pdf_page.rect.height * zoom)))
            raw = pdf_page.get_text("dict")
            blocks: list[OCRTextBlock] = []
            for block in raw.get("blocks", []):
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    spans = line.get("spans", [])
                    text = "".join(str(span.get("text", "")) for span in spans).strip()
                    if not text:
                        continue
                    bbox = line.get("bbox")
                    if not bbox and spans:
                        span_bboxes = [span["bbox"] for span in spans if "bbox" in span]
                        if not span_bboxes:
                            continue
                        bbox = (
                            min(float(item[0]) for item in span_bboxes),
                            min(float(item[1]) for item in span_bboxes),
                            max(float(item[2]) for item in span_bboxes),
                            max(float(item[3]) for item in span_bboxes),
                        )
                    if not bbox:
                        continue
                    x0, y0, x1, y1 = [float(value) * zoom for value in bbox]
                    if x1 <= x0 or y1 <= y0:
                        continue
                    blocks.append(
                        OCRTextBlock(
                            text=text,
                            polygon=[
                                [x0, y0],
                                [x1, y0],
                                [x1, y1],
                                [x0, y1],
                            ],
                            confidence=1.0,
                        )
                    )
            if cache_limit > 0:
                with self._pdf_page_text_blocks_cache_lock:
                    self._pdf_page_text_blocks_cache[cache_key] = (
                        clone_blocks(blocks),
                        page_width,
                        page_height,
                    )
                    self._pdf_page_text_blocks_cache.move_to_end(cache_key)
                    while len(self._pdf_page_text_blocks_cache) > cache_limit:
                        self._pdf_page_text_blocks_cache.popitem(last=False)
            return blocks, page_width, page_height
        finally:
            doc.close()

    async def read_image(self, file_path: str) -> bytes:
        """读取图片文件"""
        _validate_path(file_path)
        with open(file_path, "rb") as f:
            return f.read()

    async def get_image_size(self, file_path: str) -> tuple[int, int]:
        """获取图片尺寸"""
        _validate_path(file_path)
        with Image.open(file_path) as img:
            return img.size
